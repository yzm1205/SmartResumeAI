import os
import logging
import tempfile
from pathlib import Path
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def format_resume_as_dict(resume_data):
    """
    Format optimized resume data for template rendering.
    
    Args:
        resume_data (dict): Optimized resume data from AI
        
    Returns:
        dict: Formatted resume data for templates
    """
    formatted_data = {
        "basic_info": resume_data.get("basic_info", {}),
        "summary": resume_data.get("basic_info", {}).get("summary", ""),
        "experiences": resume_data.get("experiences", []),
        "educations": resume_data.get("educations", []),
        "skills": resume_data.get("skills", []),
        "certifications": resume_data.get("certifications", []),
        "projects": resume_data.get("projects", []),
        "publications": resume_data.get("publications", []),
        "achievements": resume_data.get("achievements", [])
    }
    
    return formatted_data

def generate_pdf_resume(resume_data, output_path=None):
    """
    Generate PDF resume from formatted resume data.
    
    Args:
        resume_data (dict): Formatted resume data
        output_path (str, optional): Path to save PDF file. If None, a temp file is created.
        
    Returns:
        str: Path to generated PDF file
    """
    try:
        # Format data for template
        formatted_data = format_resume_as_dict(resume_data)
        
        # Create temporary file if output path not provided
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join("data", "generated_resumes")
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"resume_{timestamp}.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Name',
            fontName='Helvetica-Bold',
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=6
        ))
        styles.add(ParagraphStyle(
            name='ContactInfo',
            fontName='Helvetica',
            fontSize=9,
            alignment=TA_CENTER,
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='SectionTitle',
            fontName='Helvetica-Bold',
            fontSize=11,
            spaceAfter=6
        ))
        styles.add(ParagraphStyle(
            name='ExperienceTitle',
            fontName='Helvetica-Bold',
            fontSize=10,
            spaceAfter=2
        ))
        styles.add(ParagraphStyle(
            name='ExperienceDetails',
            fontName='Helvetica-Oblique',
            fontSize=9,
            spaceAfter=2
        ))
        styles.add(ParagraphStyle(
            name='BulletItem',
            fontName='Helvetica',
            fontSize=9,
            leftIndent=20,
            firstLineIndent=-15,
            spaceAfter=3
        ))
        styles.add(ParagraphStyle(
            name='NormalText',
            fontName='Helvetica',
            fontSize=9,
            spaceAfter=3
        ))
        
        # Build document content
        content = []
        
        # Header with name and contact info
        basic_info = formatted_data["basic_info"]
        
        content.append(Paragraph(basic_info.get("name", "Your Name"), styles["Name"]))
        
        # Contact info line
        contact_parts = []
        if basic_info.get("email"):
            contact_parts.append(basic_info["email"])
        if basic_info.get("phone"):
            contact_parts.append(basic_info["phone"])
        if basic_info.get("linkedin"):
            contact_parts.append(basic_info["linkedin"])
        
        contact_info = " | ".join(contact_parts)
        content.append(Paragraph(contact_info, styles["ContactInfo"]))
        
        # Summary section
        if formatted_data["summary"]:
            content.append(Paragraph("PROFESSIONAL SUMMARY", styles["SectionTitle"]))
            content.append(Paragraph(formatted_data["summary"], styles["NormalText"]))
            content.append(Spacer(1, 10))
        
        # Experience section
        if formatted_data["experiences"]:
            content.append(Paragraph("PROFESSIONAL EXPERIENCE", styles["SectionTitle"]))
            
            for exp in formatted_data["experiences"]:
                # Role and company
                title_line = f"{exp.get('title', 'Role')} - {exp.get('company', 'Company')}"
                content.append(Paragraph(title_line, styles["ExperienceTitle"]))
                
                # Dates and location
                location_parts = []
                if exp.get("start_date") and exp.get("end_date"):
                    location_parts.append(f"{exp.get('start_date')} - {exp.get('end_date')}")
                if exp.get("location"):
                    location_parts.append(exp.get("location"))
                
                if location_parts:
                    location_line = " | ".join(location_parts)
                    content.append(Paragraph(location_line, styles["ExperienceDetails"]))
                
                # Achievements as bullet points
                if exp.get("description"):
                    content.append(Paragraph(exp.get("description"), styles["NormalText"]))
                
                if exp.get("achievements"):
                    achievements = exp.get("achievements", [])
                    if isinstance(achievements, str):
                        achievements = achievements.split("\n")
                    
                    for achievement in achievements:
                        if achievement.strip():
                            content.append(Paragraph(f"• {achievement.strip()}", styles["BulletItem"]))
                
                content.append(Spacer(1, 5))
            
            content.append(Spacer(1, 5))
        
        # Education section
        if formatted_data["educations"]:
            content.append(Paragraph("EDUCATION", styles["SectionTitle"]))
            
            for edu in formatted_data["educations"]:
                # Degree and institution
                degree_line = f"{edu.get('degree', 'Degree')} in {edu.get('field_of_study', 'Field of Study')}"
                content.append(Paragraph(degree_line, styles["ExperienceTitle"]))
                
                # Institution and dates
                institution_parts = []
                institution_parts.append(edu.get("institution", "Institution"))
                if edu.get("start_date") and edu.get("end_date"):
                    institution_parts.append(f"{edu.get('start_date')} - {edu.get('end_date')}")
                
                institution_line = " | ".join(institution_parts)
                content.append(Paragraph(institution_line, styles["ExperienceDetails"]))
                
                # GPA
                if edu.get("gpa"):
                    content.append(Paragraph(f"GPA: {edu.get('gpa')}", styles["NormalText"]))
                
                content.append(Spacer(1, 5))
            
            content.append(Spacer(1, 5))
        
        # Skills section
        if formatted_data["skills"]:
            content.append(Paragraph("SKILLS", styles["SectionTitle"]))
            
            # Group skills by category
            skill_categories = {}
            for skill in formatted_data["skills"]:
                category = skill.get("category", "Other")
                if category not in skill_categories:
                    skill_categories[category] = []
                skill_categories[category].append(skill.get("name", ""))
            
            # Display skills by category
            for category, skills in skill_categories.items():
                content.append(Paragraph(f"{category}: {', '.join(skills)}", styles["NormalText"]))
            
            content.append(Spacer(1, 5))
        
        # Projects section
        if formatted_data["projects"]:
            content.append(Paragraph("PROJECTS", styles["SectionTitle"]))
            
            for project in formatted_data["projects"]:
                # Project name
                project_line = project.get("name", "Project")
                if project.get("technologies"):
                    project_line += f" ({project.get('technologies')})"
                content.append(Paragraph(project_line, styles["ExperienceTitle"]))
                
                # Description
                if project.get("description"):
                    content.append(Paragraph(project.get("description"), styles["NormalText"]))
                
                content.append(Spacer(1, 5))
            
            content.append(Spacer(1, 5))
        
        # Build and save PDF
        doc.build(content)
        
        logger.info(f"Generated PDF resume at: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error generating PDF resume: {str(e)}")
        return None

def generate_docx_resume(resume_data, output_path=None):
    """
    Generate DOCX resume from formatted resume data.
    
    Args:
        resume_data (dict): Formatted resume data
        output_path (str, optional): Path to save DOCX file. If None, a temp file is created.
        
    Returns:
        str: Path to generated DOCX file
    """
    try:
        # Format data for template
        formatted_data = format_resume_as_dict(resume_data)
        
        # Create temporary file if output path not provided
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join("data", "generated_resumes")
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"resume_{timestamp}.docx")
        
        # Create DOCX document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Define styles
        styles = doc.styles
        
        # Header with name and contact info
        basic_info = formatted_data["basic_info"]
        
        # Name
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(basic_info.get("name", "Your Name"))
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if basic_info.get("email"):
            contact_parts.append(basic_info["email"])
        if basic_info.get("phone"):
            contact_parts.append(basic_info["phone"])
        if basic_info.get("linkedin"):
            contact_parts.append(basic_info["linkedin"])
        
        contact_info = " | ".join(contact_parts)
        contact_paragraph = doc.add_paragraph()
        contact_run = contact_paragraph.add_run(contact_info)
        contact_run.font.size = Pt(9)
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary section
        if formatted_data["summary"]:
            doc.add_heading("PROFESSIONAL SUMMARY", level=1).bold = True
            doc.add_paragraph(formatted_data["summary"])
        
        # Experience section
        if formatted_data["experiences"]:
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=1).bold = True
            
            for exp in formatted_data["experiences"]:
                # Role and company
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(f"{exp.get('title', 'Role')} - {exp.get('company', 'Company')}")
                title_run.bold = True
                
                # Dates and location
                location_parts = []
                if exp.get("start_date") and exp.get("end_date"):
                    location_parts.append(f"{exp.get('start_date')} - {exp.get('end_date')}")
                if exp.get("location"):
                    location_parts.append(exp.get("location"))
                
                if location_parts:
                    location_line = " | ".join(location_parts)
                    location_paragraph = doc.add_paragraph()
                    location_run = location_paragraph.add_run(location_line)
                    location_run.italic = True
                    location_run.font.size = Pt(9)
                
                # Description
                if exp.get("description"):
                    doc.add_paragraph(exp.get("description"))
                
                # Achievements as bullet points
                if exp.get("achievements"):
                    achievements = exp.get("achievements", [])
                    if isinstance(achievements, str):
                        achievements = achievements.split("\n")
                    
                    for achievement in achievements:
                        if achievement.strip():
                            doc.add_paragraph(achievement.strip(), style='List Bullet')
        
        # Education section
        if formatted_data["educations"]:
            doc.add_heading("EDUCATION", level=1).bold = True
            
            for edu in formatted_data["educations"]:
                # Degree and institution
                degree_paragraph = doc.add_paragraph()
                degree_run = degree_paragraph.add_run(f"{edu.get('degree', 'Degree')} in {edu.get('field_of_study', 'Field of Study')}")
                degree_run.bold = True
                
                # Institution and dates
                institution_parts = []
                institution_parts.append(edu.get("institution", "Institution"))
                if edu.get("start_date") and edu.get("end_date"):
                    institution_parts.append(f"{edu.get('start_date')} - {edu.get('end_date')}")
                
                institution_line = " | ".join(institution_parts)
                institution_paragraph = doc.add_paragraph()
                institution_run = institution_paragraph.add_run(institution_line)
                institution_run.italic = True
                institution_run.font.size = Pt(9)
                
                # GPA
                if edu.get("gpa"):
                    doc.add_paragraph(f"GPA: {edu.get('gpa')}")
        
        # Skills section
        if formatted_data["skills"]:
            doc.add_heading("SKILLS", level=1).bold = True
            
            # Group skills by category
            skill_categories = {}
            for skill in formatted_data["skills"]:
                category = skill.get("category", "Other")
                if category not in skill_categories:
                    skill_categories[category] = []
                skill_categories[category].append(skill.get("name", ""))
            
            # Display skills by category
            for category, skills in skill_categories.items():
                doc.add_paragraph(f"{category}: {', '.join(skills)}")
        
        # Projects section
        if formatted_data["projects"]:
            doc.add_heading("PROJECTS", level=1).bold = True
            
            for project in formatted_data["projects"]:
                # Project name
                project_paragraph = doc.add_paragraph()
                project_text = project.get("name", "Project")
                if project.get("technologies"):
                    project_text += f" ({project.get('technologies')})"
                project_run = project_paragraph.add_run(project_text)
                project_run.bold = True
                
                # Description
                if project.get("description"):
                    doc.add_paragraph(project.get("description"))
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Generated DOCX resume at: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error generating DOCX resume: {str(e)}")
        return None 